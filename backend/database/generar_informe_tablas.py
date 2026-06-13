"""Genera documento Word con las 14 tablas de agri_junin."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "Informe_14_Tablas_AgriJunin.docx"

TABLES = [
    {
        "num": 1,
        "name": "roles",
        "type": "Catálogo",
        "desc": "Roles de acceso del sistema (administrador, agricultor, técnico).",
        "fields": [
            ("id", "TINYINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del rol"),
            ("codigo", "VARCHAR(20)", "NO", "UNIQUE", "Código interno: administrador, agricultor, tecnico"),
            ("nombre", "VARCHAR(60)", "NO", "—", "Nombre legible del rol"),
            ("descripcion", "VARCHAR(255)", "SÍ", "—", "Descripción de funciones del rol"),
        ],
        "refs": "Referenciada por: usuarios (rol_id)",
    },
    {
        "num": 2,
        "name": "tipos_cultivo",
        "type": "Catálogo",
        "desc": "Clasificación de cultivos (cereal, tubérculo, hortaliza, legumbre, etc.).",
        "fields": [
            ("id", "TINYINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador"),
            ("codigo", "VARCHAR(30)", "NO", "UNIQUE", "Código del tipo de cultivo"),
            ("nombre", "VARCHAR(80)", "NO", "—", "Nombre del tipo"),
            ("descripcion", "VARCHAR(255)", "SÍ", "—", "Ejemplos de cultivos del tipo"),
        ],
        "refs": "Referenciada por: cultivos (tipo_id)",
    },
    {
        "num": 3,
        "name": "temporadas_cultivo",
        "type": "Catálogo",
        "desc": "Época de siembra o cultivo (verano, invierno, todo el año).",
        "fields": [
            ("id", "TINYINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador"),
            ("codigo", "VARCHAR(20)", "NO", "UNIQUE", "Código: verano, invierno, todo_año"),
            ("nombre", "VARCHAR(60)", "NO", "—", "Nombre legible de la temporada"),
        ],
        "refs": "Referenciada por: cultivos (temporada_id)",
    },
    {
        "num": 4,
        "name": "perfiles_climaticos",
        "type": "Catálogo",
        "desc": "Rangos óptimos de humedad y temperatura para generar alertas.",
        "fields": [
            ("id", "SMALLINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador"),
            ("nombre", "VARCHAR(100)", "NO", "—", "Nombre del perfil (ej. Perfil Papa)"),
            ("humedad_optima_min", "DECIMAL(5,2)", "SÍ", "—", "Humedad mínima ideal (%)"),
            ("humedad_optima_max", "DECIMAL(5,2)", "SÍ", "—", "Humedad máxima ideal (%)"),
            ("temp_optima_min", "DECIMAL(5,2)", "SÍ", "—", "Temperatura mínima ideal (°C)"),
            ("temp_optima_max", "DECIMAL(5,2)", "SÍ", "—", "Temperatura máxima ideal (°C)"),
            ("descripcion", "VARCHAR(255)", "SÍ", "—", "Descripción del perfil climático"),
        ],
        "refs": "Referenciada por: cultivos (perfil_climatico_id)",
    },
    {
        "num": 5,
        "name": "tipos_suelo",
        "type": "Catálogo",
        "desc": "Textura del suelo de los lotes agrícolas.",
        "fields": [
            ("id", "TINYINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador"),
            ("codigo", "VARCHAR(20)", "NO", "UNIQUE", "Código: arcilloso, franco, arenoso, limoso"),
            ("nombre", "VARCHAR(60)", "NO", "—", "Nombre del tipo de suelo"),
            ("descripcion", "VARCHAR(255)", "SÍ", "—", "Características del suelo"),
        ],
        "refs": "Referenciada por: lotes (tipo_suelo_id)",
    },
    {
        "num": 6,
        "name": "tipos_sensor",
        "type": "Catálogo",
        "desc": "Tipos de dispositivos IoT instalados en los lotes.",
        "fields": [
            ("id", "TINYINT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador"),
            ("codigo", "VARCHAR(30)", "NO", "UNIQUE", "Código: humedad_suelo, temperatura, ph, etc."),
            ("nombre", "VARCHAR(80)", "NO", "—", "Nombre del tipo de sensor"),
            ("unidad_medida_def", "VARCHAR(20)", "NO", "DEFAULT '%'", "Unidad por defecto: %, °C, pH, mm, lux"),
            ("descripcion", "VARCHAR(255)", "SÍ", "—", "Descripción de qué mide el sensor"),
        ],
        "refs": "Referenciada por: sensores (tipo_sensor_id)",
    },
    {
        "num": 7,
        "name": "usuarios",
        "type": "Entidad",
        "desc": "Identidad y autenticación de usuarios del sistema.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del usuario"),
            ("nombre", "VARCHAR(120)", "NO", "—", "Nombre completo"),
            ("email", "VARCHAR(180)", "NO", "UNIQUE", "Correo electrónico (login)"),
            ("dni", "VARCHAR(8)", "SÍ", "UNIQUE", "DNI peruano (login alternativo)"),
            ("password", "VARCHAR(255)", "NO", "—", "Contraseña encriptada (bcrypt)"),
            ("rol_id", "TINYINT UNSIGNED", "NO", "FK → roles, DEFAULT 2", "Rol del usuario"),
            ("activo", "TINYINT(1)", "NO", "DEFAULT 1", "1=activo, 0=inactivo"),
            ("estado_cuenta", "ENUM", "NO", "aprobada/pendiente/rechazada", "Estado de aprobación de cuenta"),
            ("avatar_url", "VARCHAR(255)", "SÍ", "—", "URL de imagen de perfil"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: rol_id → roles | Referenciada por: agricultores, registros_agricolas, solicitudes_aprobacion",
    },
    {
        "num": 8,
        "name": "agricultores",
        "type": "Entidad",
        "desc": "Extensión 1:1 del usuario con rol agricultor. Datos del dominio agrícola.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del agricultor"),
            ("usuario_id", "INT UNSIGNED", "NO", "FK → usuarios, UNIQUE", "Relación 1:1 con usuario"),
            ("telefono", "VARCHAR(20)", "SÍ", "—", "Teléfono de contacto"),
            ("direccion", "TEXT", "SÍ", "—", "Dirección del agricultor"),
            ("distrito", "VARCHAR(80)", "NO", "DEFAULT 'Junín'", "Distrito"),
            ("provincia", "VARCHAR(80)", "NO", "DEFAULT 'Junín'", "Provincia"),
            ("departamento", "VARCHAR(80)", "NO", "DEFAULT 'Junín'", "Departamento"),
            ("hectareas_totales", "DECIMAL(10,2)", "NO", "DEFAULT 0", "Superficie total en hectáreas"),
            ("fecha_registro", "DATE", "NO", "—", "Fecha de alta como agricultor"),
            ("activo", "TINYINT(1)", "NO", "DEFAULT 1", "Estado del agricultor"),
            ("notas", "TEXT", "SÍ", "—", "Observaciones adicionales"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: usuario_id → usuarios | Referenciada por: lotes (agricultor_id)",
    },
    {
        "num": 9,
        "name": "cultivos",
        "type": "Entidad",
        "desc": "Catálogo de cultivos gestionados en el sistema.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del cultivo"),
            ("nombre", "VARCHAR(120)", "NO", "INDEX", "Nombre común (ej. Papa, Maíz)"),
            ("nombre_cientifico", "VARCHAR(150)", "SÍ", "—", "Nombre científico"),
            ("tipo_id", "TINYINT UNSIGNED", "NO", "FK → tipos_cultivo", "Tipo de cultivo"),
            ("temporada_id", "TINYINT UNSIGNED", "NO", "FK → temporadas_cultivo", "Temporada de siembra"),
            ("perfil_climatico_id", "SMALLINT UNSIGNED", "SÍ", "FK → perfiles_climaticos", "Perfil climático óptimo"),
            ("dias_crecimiento", "SMALLINT UNSIGNED", "NO", "DEFAULT 90", "Días del ciclo de crecimiento"),
            ("rendimiento_promedio", "DECIMAL(8,2)", "SÍ", "—", "Rendimiento promedio (qq/ha)"),
            ("descripcion", "TEXT", "SÍ", "—", "Descripción del cultivo"),
            ("activo", "TINYINT(1)", "NO", "DEFAULT 1", "Visible/activo en el sistema"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: tipo_id, temporada_id, perfil_climatico_id | Referenciada por: lotes (cultivo_id)",
    },
    {
        "num": 10,
        "name": "lotes",
        "type": "Entidad (Central)",
        "desc": "Tabla central del sistema. Parcelas agrícolas de cada agricultor.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del lote"),
            ("agricultor_id", "INT UNSIGNED", "NO", "FK → agricultores", "Dueño del lote"),
            ("cultivo_id", "INT UNSIGNED", "SÍ", "FK → cultivos", "Cultivo asignado"),
            ("codigo_lote", "VARCHAR(30)", "NO", "UNIQUE", "Código único (ej. LOT-CHU-001)"),
            ("nombre", "VARCHAR(120)", "NO", "—", "Nombre del lote"),
            ("ubicacion", "VARCHAR(255)", "SÍ", "—", "Descripción de ubicación"),
            ("latitud", "DECIMAL(10,7)", "SÍ", "—", "Coordenada GPS latitud"),
            ("longitud", "DECIMAL(10,7)", "SÍ", "—", "Coordenada GPS longitud"),
            ("area_hectareas", "DECIMAL(10,2)", "NO", "—", "Área en hectáreas"),
            ("tipo_suelo_id", "TINYINT UNSIGNED", "NO", "FK → tipos_suelo, DEFAULT 3", "Tipo de suelo"),
            ("estado", "ENUM", "NO", "preparacion/siembra/crecimiento/cosecha/barbecho", "Fase del lote"),
            ("fecha_siembra", "DATE", "SÍ", "—", "Fecha de siembra"),
            ("fecha_cosecha_est", "DATE", "SÍ", "—", "Fecha estimada de cosecha"),
            ("activo", "TINYINT(1)", "NO", "DEFAULT 1", "Requiere aprobación si es 0"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: agricultor_id, cultivo_id, tipo_suelo_id | Referenciada por: sensores, registros_agricolas, solicitudes_aprobacion",
    },
    {
        "num": 11,
        "name": "solicitudes_aprobacion",
        "type": "Entidad",
        "desc": "Flujo de aprobación de lotes y cultivos creados por agricultores.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador de la solicitud"),
            ("entidad_tipo", "ENUM", "NO", "lote / cultivo", "Tipo de entidad solicitada"),
            ("entidad_id", "INT UNSIGNED", "NO", "INDEX", "ID del lote o cultivo"),
            ("estado", "ENUM", "NO", "pendiente/aprobado/rechazado", "Estado de la solicitud"),
            ("solicitado_por", "INT UNSIGNED", "NO", "FK → usuarios", "Agricultor solicitante"),
            ("revisado_por", "INT UNSIGNED", "SÍ", "FK → usuarios", "Admin o técnico revisor"),
            ("fecha_solicitud", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de la solicitud"),
            ("fecha_revision", "TIMESTAMP", "SÍ", "—", "Fecha de revisión"),
            ("motivo_rechazo", "VARCHAR(255)", "SÍ", "—", "Motivo si fue rechazada"),
            ("lote_destino_id", "INT UNSIGNED", "SÍ", "FK → lotes", "Lote destino (cultivos nuevos)"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: solicitado_por, revisado_por → usuarios | lote_destino_id → lotes",
    },
    {
        "num": 12,
        "name": "sensores",
        "type": "Entidad",
        "desc": "Dispositivos IoT instalados en lotes para monitoreo en tiempo real.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del sensor"),
            ("lote_id", "INT UNSIGNED", "NO", "FK → lotes", "Lote donde está instalado"),
            ("tipo_sensor_id", "TINYINT UNSIGNED", "NO", "FK → tipos_sensor", "Tipo de sensor"),
            ("codigo_sensor", "VARCHAR(40)", "NO", "UNIQUE", "Código único del dispositivo"),
            ("nombre", "VARCHAR(120)", "NO", "—", "Nombre descriptivo"),
            ("unidad_medida", "VARCHAR(20)", "NO", "—", "Unidad: %, °C, pH, mm, lux"),
            ("valor_min", "DECIMAL(10,2)", "SÍ", "—", "Valor mínimo operativo"),
            ("valor_max", "DECIMAL(10,2)", "SÍ", "—", "Valor máximo operativo"),
            ("ultima_lectura", "DECIMAL(12,4)", "SÍ", "—", "Último valor medido"),
            ("ultima_fecha", "DATETIME", "SÍ", "—", "Fecha de última lectura"),
            ("estado", "ENUM", "NO", "activo/inactivo/mantenimiento/falla", "Estado del sensor"),
            ("bateria_pct", "TINYINT UNSIGNED", "SÍ", "0–100", "Porcentaje de batería"),
            ("activo", "TINYINT(1)", "NO", "DEFAULT 1", "Activo en el sistema"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: lote_id → lotes | tipo_sensor_id → tipos_sensor | Referenciada por: alertas (sensor_id)",
    },
    {
        "num": 13,
        "name": "registros_agricolas",
        "type": "Entidad",
        "desc": "Mediciones de campo registradas por técnicos (temperatura, humedad, pH, producción).",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador del registro"),
            ("lote_id", "INT UNSIGNED", "NO", "FK → lotes", "Lote al que pertenece"),
            ("fecha_registro", "DATETIME", "NO", "DEFAULT NOW, INDEX", "Fecha y hora del registro"),
            ("temperatura", "DECIMAL(6,2)", "SÍ", "—", "Temperatura en °C"),
            ("humedad_suelo", "DECIMAL(6,2)", "SÍ", "—", "Humedad del suelo (%)"),
            ("humedad_aire", "DECIMAL(6,2)", "SÍ", "—", "Humedad del aire (%)"),
            ("ph_suelo", "DECIMAL(4,2)", "SÍ", "—", "pH del suelo (0–14)"),
            ("precipitacion_mm", "DECIMAL(8,2)", "SÍ", "DEFAULT 0", "Precipitación en mm"),
            ("produccion_kg", "DECIMAL(12,2)", "SÍ", "—", "Producción cosechada en kg"),
            ("observaciones", "TEXT", "SÍ", "—", "Notas del técnico"),
            ("registrado_por", "INT UNSIGNED", "SÍ", "FK → usuarios", "Usuario que registró"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: lote_id → lotes | registrado_por → usuarios | Referenciada por: alertas (registro_id)",
    },
    {
        "num": 14,
        "name": "alertas",
        "type": "Entidad",
        "desc": "Alertas generadas por condiciones anómalas en registros o sensores.",
        "fields": [
            ("id", "INT UNSIGNED", "NO", "PK, AUTO_INCREMENT", "Identificador de la alerta"),
            ("registro_id", "INT UNSIGNED", "SÍ", "FK → registros_agricolas", "Origen: registro agrícola"),
            ("sensor_id", "INT UNSIGNED", "SÍ", "FK → sensores", "Origen: sensor IoT"),
            ("tipo", "ENUM", "NO", "humedad/temperatura/ph/pluvia/sensor/produccion/sistema", "Tipo de alerta"),
            ("nivel", "ENUM", "NO", "info/advertencia/critica", "Nivel de gravedad"),
            ("titulo", "VARCHAR(200)", "NO", "—", "Título de la alerta"),
            ("mensaje", "TEXT", "NO", "—", "Descripción detallada"),
            ("leida", "TINYINT(1)", "NO", "DEFAULT 0", "Si fue leída por el usuario"),
            ("resuelta", "TINYINT(1)", "NO", "DEFAULT 0", "Si fue resuelta"),
            ("fecha_alerta", "DATETIME", "NO", "DEFAULT NOW", "Cuándo se generó"),
            ("fecha_resolucion", "DATETIME", "SÍ", "—", "Cuándo se resolvió"),
            ("created_at", "TIMESTAMP", "NO", "DEFAULT NOW", "Fecha de creación"),
            ("updated_at", "TIMESTAMP", "NO", "ON UPDATE NOW", "Última modificación"),
        ],
        "refs": "FK: registro_id → registros_agricolas | sensor_id → sensores | CHECK: origen obligatorio",
    },
]


def add_table_section(doc: Document, table_info: dict) -> None:
    heading = doc.add_heading(
        f"Tabla {table_info['num']}: {table_info['name']} ({table_info['type']})",
        level=2,
    )
    heading.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

    p = doc.add_paragraph(table_info["desc"])
    p.paragraph_format.space_after = Pt(6)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    headers = ["Campo", "Tipo de dato", "Nulo", "Restricción", "Descripción"]
    for i, text in enumerate(headers):
        hdr[i].text = text
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for field in table_info["fields"]:
        row = tbl.add_row().cells
        for i, val in enumerate(field):
            row[i].text = val

    for row in tbl.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph(table_info["refs"]).runs[0].italic = True
    doc.add_paragraph()


def main() -> None:
    doc = Document()

    title = doc.add_heading(
        "Base de Datos — Sistema AgriJunín\nAgricultura Inteligente Región Junín",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Documento técnico con la descripción de las 14 tablas de la base de datos "
        "agri_junin (MySQL/MariaDB, normalización 3FN, motor InnoDB, codificación utf8mb4)."
    )

    doc.add_heading("Resumen", level=1)
    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    sh = summary.rows[0].cells
    for i, t in enumerate(["#", "Tabla", "Tipo", "Campos"]):
        sh[i].text = t
        sh[i].paragraphs[0].runs[0].bold = True
    for t in TABLES:
        r = summary.add_row().cells
        r[0].text = str(t["num"])
        r[1].text = t["name"]
        r[2].text = t["type"]
        r[3].text = str(len(t["fields"]))

    doc.add_paragraph()
    doc.add_heading("Catálogos (6 tablas)", level=1)
    doc.add_paragraph(
        "Tablas de referencia reutilizables que eliminan redundancia y garantizan integridad (3FN)."
    )

    doc.add_heading("Entidades (8 tablas)", level=1)
    doc.add_paragraph(
        "Tablas transaccionales que almacenan la información operativa del sistema."
    )

    doc.add_page_break()
    doc.add_heading("Descripción detallada de las 14 tablas", level=1)

    for table_info in TABLES:
        add_table_section(doc, table_info)

    doc.add_heading("Información general", level=1)
    info = [
        ("Motor de BD", "MySQL / MariaDB (InnoDB)"),
        ("Nombre de BD", "agri_junin"),
        ("Codificación", "utf8mb4_unicode_ci"),
        ("Normalización", "Tercera Forma Normal (3FN)"),
        ("Total tablas", "14 (6 catálogos + 8 entidades)"),
        ("Procedimientos almacenados", "18"),
        ("Archivo schema", "backend/database/schema.sql"),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].text = v

    doc.save(OUTPUT)
    print(f"[OK] Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
